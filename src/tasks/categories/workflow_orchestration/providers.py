"""
Consolidated providers for workflow orchestration category.
Handles file placement, directory creation, setup configuration, and evaluation logic.
"""

import os
from typing import Dict, Any, List, Optional
from PIL import Image, ImageDraw
from docx import Document

from ..base import (
    FileProviderInterface,
    ConfigProviderInterface,
    EvaluationProviderInterface,
)
from .evaluators import WorkflowOrchestrationEvaluators


class WorkflowOrchestrationFileProvider(FileProviderInterface):
    """File provider implementation for workflow orchestration tasks."""

    def __init__(self):
        """Initialize the file provider."""
        self.supported_tasks = {"basic_info_gathering"}
        self.directory_creation_mapping = {"basic_info_gathering": ["/home/user/Desktop"]}

    def get_file_placement_path(self, task_type: str, filename: str) -> str:
        """Get the file placement path for a specific task type."""
        placement_mapping = {"basic_info_gathering": "/home/user/Desktop"}
        base_path = placement_mapping.get(task_type, "/home/user/Desktop")
        return os.path.join(base_path, filename)

    def get_directories_to_create(self, task_type: str) -> List[str]:
        """Get list of directories that need to be created for a task type."""
        return self.directory_creation_mapping.get(task_type, [])

    def create_task_files(self, task_data: Dict[str, Any], task_id: str, temp_dir: str) -> Optional[Dict[str, str]]:
        """Create workflow orchestration specific task files with proper structure."""
        task_type = task_data.get("task_type")
        if not self.supports_task_type(task_type):
            return None
        files_created = {}
        if task_type == "basic_info_gathering":
            research_filename = task_data.get("research_filename", "research.html")
            research_content = task_data.get("research_html_content", "<html><body>Research Content</body></html>")
            research_path = os.path.join(temp_dir, research_filename)
            with open(research_path, "w", encoding="utf-8") as f:
                f.write(research_content)
            files_created["main_file"] = research_path
            files_created["main_filename"] = research_filename
            source_image_filename = task_data.get("source_image_filename", "source_image.png")
            source_image_path = self._generate_source_image(task_data, temp_dir, source_image_filename)
            files_created["source_image"] = source_image_path
            additional_files = {
                source_image_filename: {
                    "local_path": source_image_path,
                    "filename": source_image_filename,
                }
            }
            files_created["additional_files"] = additional_files
            expected_files = self._create_expected_files(task_data, task_id, temp_dir)
            if expected_files:
                files_created.update(expected_files)
        return files_created

    def supports_task_type(self, task_type: str) -> bool:
        """Check if this provider supports the given task type."""
        return task_type in self.supported_tasks

    def _generate_source_image(self, task_data: Dict[str, Any], temp_dir: str, filename: str) -> str:
        """Generate source image for workflow tasks."""
        size = task_data.get("image_size", (800, 600))
        img = Image.new("RGB", size, color="lightblue")
        draw = ImageDraw.Draw(img)
        draw.rectangle([50, 50, size[0] - 50, size[1] - 50], fill="white", outline="black", width=3)
        image_path = os.path.join(temp_dir, filename)
        img.save(image_path, "PNG")
        return image_path

    def _create_expected_files(self, task_data: Dict[str, Any], task_id: str, temp_dir: str) -> Dict[str, str]:
        """Create expected files for evaluation."""
        expected_files = {}
        expected_doc_path = os.path.join(temp_dir, "expected_report.docx")
        doc = Document()
        doc.add_heading("Expected Research Report", 0)
        doc.add_paragraph(task_data.get("expected_summary", "Expected content summary"))
        doc.save(expected_doc_path)
        expected_files["expected_document"] = expected_doc_path
        json_spec = {"global": {"required_texts": [task_data.get("expected_summary", "Expected summary")], "required_headings": []}}
        import json

        json_spec_path = os.path.join(temp_dir, "expected_spec.json")
        with open(json_spec_path, "w", encoding="utf-8") as f:
            json.dump(json_spec, f, indent=2)
        expected_files["json_specification"] = json_spec_path
        return expected_files


class WorkflowOrchestrationConfigProvider(ConfigProviderInterface):
    """Config provider implementation for workflow orchestration tasks."""

    def __init__(self):
        """Initialize the config provider."""
        self.supported_tasks = {"basic_info_gathering"}

    def build_setup_steps(
        self,
        task_data: Dict[str, Any],
        s3_urls: Dict[str, str],
        files_created: Dict[str, str],
    ) -> List[Dict[str, Any]]:
        """Build setup configuration steps for task initialization."""
        task_type = task_data.get("task_type")
        return self.build_config(task_type, task_data, files_created, s3_urls)

    def build_config(
        self,
        task_type: str,
        task_data: Dict[str, Any],
        files_created: Dict[str, str],
        s3_urls: Dict[str, str] = None,
    ) -> List[Dict[str, Any]]:
        """Build configuration for workflow orchestration tasks."""
        if not self.supports_task_type(task_type):
            raise ValueError(f"Unsupported task type: {task_type}")
        if task_type == "basic_info_gathering":
            return self._build_basic_info_gathering_config(task_data, files_created, s3_urls)
        return []

    def _build_basic_info_gathering_config(
        self,
        task_data: Dict[str, Any],
        files_created: Dict[str, str],
        s3_urls: Dict[str, str] = None,
    ) -> List[Dict[str, Any]]:
        """Build config for basic information gathering task."""
        research_filename = task_data.get("research_filename", "research.html")
        source_image_filename = task_data.get("source_image_filename", "source_image.png")
        download_files: List[Dict[str, str]] = []
        if s3_urls and "main_file" in s3_urls:
            download_files.append({"url": s3_urls["main_file"], "path": f"/tmp/{research_filename}"})
        else:
            download_files.append({"url": f"file://{files_created['main_file']}", "path": f"/tmp/{research_filename}"})
        additional_entries_added = False
        if s3_urls and "additional_files" in s3_urls:
            for _orig, file_info in s3_urls["additional_files"].items():
                download_files.append({"url": file_info["url"], "path": f"/home/user/Desktop/{file_info['filename']}"})
                additional_entries_added = True
        if not additional_entries_added:
            if s3_urls and "source_image" in s3_urls:
                download_files.append({"url": s3_urls["source_image"], "path": f"/home/user/Desktop/{source_image_filename}"})
            else:
                download_files.append({"url": f"file://{files_created['source_image']}", "path": f"/home/user/Desktop/{source_image_filename}"})
        config: List[Dict[str, Any]] = [
            {"type": "download", "parameters": {"files": download_files}},
            {
                "type": "execute",
                "parameters": {
                    "command": [
                        "python3",
                        "-c",
                        "import subprocess, os; log=open('/tmp/http_server.log','a'); "
                        "p=subprocess.Popen(['python3','-m','http.server','8080','--directory','/tmp'], stdout=log, stderr=log, preexec_fn=os.setsid); "
                        "open('/tmp/http_server.pid','w').write(str(p.pid)); print('Server started on port 8080')",
                    ]
                },
            },
            {"type": "sleep", "parameters": {"seconds": 2}},
            {
                "type": "launch",
                "parameters": {
                    "command": [
                        "google-chrome",
                        "--new-window",
                        f"http://localhost:8080/{research_filename}",
                    ]
                },
            },
            {"type": "sleep", "parameters": {"seconds": 2}},
            {"type": "launch", "parameters": {"command": ["libreoffice", "--writer"]}},
            {"type": "sleep", "parameters": {"seconds": 2}},
            {
                "type": "launch",
                "parameters": {"command": ["gimp", f"/home/user/Desktop/{source_image_filename}"]},
            },
            {"type": "sleep", "parameters": {"seconds": 3}},
            {"type": "execute", "parameters": {"command": ["/bin/bash", "-lc", "wmctrl -k on"]}},
            {"type": "sleep", "parameters": {"seconds": 10.0}},
        ]
        return config

    def get_evaluation_mode(self, task_type: str, level: int) -> str:
        """Get the evaluation mode for a specific task type and level."""
        return "multi_evaluator"

    def supports_task_type(self, task_type: str) -> bool:
        """Check if this provider supports the given task type."""
        return task_type in self.supported_tasks


class WorkflowOrchestrationEvaluationProvider(EvaluationProviderInterface):
    """Evaluation provider implementation for workflow orchestration tasks."""

    def __init__(self):
        """Initialize the evaluation provider."""
        self.supported_tasks = {"basic_info_gathering"}

    def build_evaluator_config(
        self,
        task_data: Dict[str, Any],
        files_created: Dict[str, str],
        evaluation_mode: str,
        s3_urls: Dict[str, str],
    ) -> Dict[str, Any]:
        """Build evaluator configuration for the task."""
        task_type = task_data.get("task_type")
        return self.build_evaluation_config(task_type, task_data, files_created, s3_urls)

    def build_evaluation_config(
        self,
        task_type: str,
        task_data: Dict[str, Any],
        files_created: Dict[str, str],
        s3_urls: Dict[str, str] = None,
    ) -> Dict[str, Any]:
        """Build evaluation configuration for workflow orchestration tasks."""
        if not self.supports_task_type(task_type):
            raise ValueError(f"Unsupported task type: {task_type}")
        if task_type == "basic_info_gathering":
            return self._build_basic_info_gathering_evaluation(task_data, files_created, s3_urls)
        return {}

    def _build_basic_info_gathering_evaluation(
        self,
        task_data: Dict[str, Any],
        files_created: Dict[str, str],
        s3_urls: Dict[str, str] = None,
    ) -> Dict[str, Any]:
        """Build evaluation config for basic information gathering task."""
        expected_document_url = s3_urls.get("expected_document", "") if s3_urls else ""
        evaluation_config = {
            "postconfig": [
                {"type": "activate_window", "parameters": {"window_name": "LibreOffice Writer", "strict": False}},
                {"type": "execute", "parameters": {"command": ["python3", "-c", "import pyautogui; pyautogui.hotkey('ctrl', 's')"]}},
                {"type": "sleep", "parameters": {"seconds": 2}},
            ],
            "func": ["compare_docx_images", "compare_answer", "check_image_size", "evaluate_document_against_spec"],
            "result": [
                {
                    "type": "vm_file",
                    "path": task_data.get("expected_files", {}).get("document_path", "/home/user/Desktop/" + task_data.get("output_document", "report.docx")),
                    "dest": "result_document.docx",
                },
                {"type": "rule", "rules": {"answer": "Exists"}},
                {
                    "type": "vm_file",
                    "path": task_data.get("expected_files", {}).get("image_path", "/home/user/Desktop/" + task_data.get("processed_image", "processed_image.png")),
                    "dest": "result_image.png",
                },
                {
                    "type": "vm_file",
                    "path": task_data.get("expected_files", {}).get("document_path", "/home/user/Desktop/" + task_data.get("output_document", "report.docx")),
                    "dest": "eval_document.docx",
                },
            ],
            "expected": [
                {"type": "cloud_file", "path": expected_document_url, "dest": "expected_document.docx"},
                {
                    "type": "vm_command_line",
                    "command": f"test -f {task_data.get('expected_files', {}).get('image_path', '/home/user/Desktop/' + task_data.get('processed_image', 'processed_image.png'))} && echo 'Exists' || echo 'Missing'",
                },
                {"type": "rule", "rules": {"width": task_data.get("target_width", 800), "height": task_data.get("target_height", 600)}},
                {"type": "cloud_file", "path": s3_urls.get("json_specification", "") if s3_urls else "", "dest": "expected_spec.json"},
            ],
            "options": [{"debug": True}, {}, {}, {"debug": True, "threshold": 0.75}],
            "conj": "and",
        }
        return evaluation_config

    def supports_task_type(self, task_type: str) -> bool:
        """Check if this provider supports the given task type."""
        return task_type in self.supported_tasks

    def get_evaluator_instance(self):
        """Get the evaluator instance for this category."""
        return WorkflowOrchestrationEvaluators()


__all__ = [
    "WorkflowOrchestrationFileProvider",
    "WorkflowOrchestrationConfigProvider",
    "WorkflowOrchestrationEvaluationProvider",
]
