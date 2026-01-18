"""
Consolidated providers for image processing category.
Handles file placement, directory creation, setup configuration, and evaluation logic.
"""

import os
import json
from typing import Dict, List, Any, Optional, Tuple
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches

from ..base import (
    FileProviderInterface,
    ConfigProviderInterface,
    EvaluationProviderInterface,
)
from .evaluators import ImageProcessingEvaluators


class ImageProcessingFileProvider(FileProviderInterface):
    """File provider implementation for image processing tasks."""

    def __init__(self):
        """Initialize image processing file provider."""
        self.supported_tasks = {
            "basic_image_insertion",
            "image_resize_insertion",
            "image_modify_caption",
        }

        # File placement paths
        self.file_placement_mapping = {
            "basic_image_insertion": "/home/user/Desktop/{filename}",
            "image_resize_insertion": "/home/user/Desktop/{filename}",
            "image_modify_caption": "/home/user/Desktop/{filename}",
        }

        # Directory creation mapping (Desktop already exists in Ubuntu)
        self.directory_creation_mapping = {
            "basic_image_insertion": [],
            "image_resize_insertion": [],
            "image_modify_caption": [],
        }

    def get_file_placement_path(self, task_type: str, filename: str) -> str:
        """Get the file placement path for a specific task type."""
        if not self.supports_task_type(task_type):
            raise ValueError(f"Unsupported task type: {task_type}")

        path_template = self.file_placement_mapping[task_type]
        return path_template.format(filename=filename)

    def get_directories_to_create(self, task_type: str) -> List[str]:
        """Get list of directories that need to be created for a task type."""
        if not self.supports_task_type(task_type):
            return []
        return self.directory_creation_mapping.get(task_type, [])

    def create_task_files(self, task_data: Dict[str, Any], task_id: str, temp_dir: str) -> Optional[Dict[str, str]]:
        """Create image processing specific task files."""
        task_type = task_data.get("task_type")

        if not self.supports_task_type(task_type):
            return None

        files_created = {}

        # Generate source image
        source_image_path = self._generate_source_image(task_data, temp_dir)
        files_created["main_file"] = source_image_path  # Main file for file manager
        files_created["source_image"] = source_image_path  # For our internal use
        files_created["main_filename"] = task_data.get("source_image", "source_image.png")

        # Generate document template
        document_path = self._generate_document_template(task_data, temp_dir)
        files_created["document_template"] = document_path

        # Ensure document template is included in outputs by adding to additional_files
        document_filename = task_data.get("document_file", "document.docx")
        additional_files: Dict[str, Dict[str, str]] = {}
        additional_files[document_filename] = {
            "local_path": document_path,
            "filename": document_filename,
        }
        files_created["additional_files"] = additional_files

        # Create expected/ground truth files using the ACTUAL source files
        expected_files = self._create_expected_files(task_data, task_id, temp_dir, source_image_path, document_path)
        if expected_files:
            files_created.update(expected_files)

        # For level 3 tasks, include JSON specification in additional files
        if task_type == "image_modify_caption" and "json_specification" in expected_files:
            spec_path = expected_files["json_specification"]
            base_name = document_filename.replace(".docx", "")
            spec_filename = f"{base_name}_spec.json"

            # Add JSON spec to additional files for local output
            additional_files[spec_filename] = {
                "local_path": spec_path,
                "filename": spec_filename,
            }
            files_created["additional_files"] = additional_files

        return files_created

    def supports_task_type(self, task_type: str) -> bool:
        """Check if this provider supports the given task type."""
        return task_type in self.supported_tasks

    def _generate_source_image(self, task_data: Dict[str, Any], temp_dir: str) -> str:
        """Generate professional source image based on enhanced task requirements."""
        # Get professional context from enhanced task data
        professional_scenario = task_data.get("professional_scenario", {})
        domain_context = task_data.get("domain_context", {})

        # Get image size based on task type and professional requirements
        if task_data.get("task_type") == "image_resizing_placement":
            size = task_data.get("source_image_size", (1600, 1200))
        else:
            size = task_data.get("image_size", (1200, 800))

        # Create professional image based on domain and scenario
        if professional_scenario:
            img = self._create_professional_image(professional_scenario, domain_context, size, task_data)
        else:
            # Fallback to enhanced image generation for backward compatibility
            image_type = task_data.get("image_type", "corporate_presentation")
            img = self._create_enhanced_image_by_type(image_type, size)

        # Save image with professional quality
        source_filename = task_data.get("source_image", "source_image.png")
        image_path = os.path.join(temp_dir, source_filename)

        # Determine format and save with high quality
        if source_filename.lower().endswith(".jpg") or source_filename.lower().endswith(".jpeg"):
            img.save(image_path, "JPEG", quality=98, optimize=True)
        elif source_filename.lower().endswith(".png"):
            img.save(image_path, "PNG", optimize=True)
        elif source_filename.lower().endswith(".bmp"):
            img.save(image_path, "BMP")
        elif source_filename.lower().endswith(".tiff"):
            img.save(image_path, "TIFF", compression="lzw")
        else:
            img.save(image_path, "PNG", optimize=True)

        return image_path

    def _create_professional_image(self, professional_scenario: Dict[str, Any], domain_context: Dict[str, Any], size: Tuple[int, int], task_data: Dict[str, Any]) -> Image.Image:
        """Create professional-quality image based on domain and scenario."""
        domain = domain_context.get("domain", "business_presentation")
        professional_type = professional_scenario.get("professional_type", "corporate_presentation")
        layout = professional_scenario.get("layout", "slide_format")
        color_scheme = professional_scenario.get("color_scheme", "#3498db")

        # Generate unique variations based on task seed
        seed = task_data.get("seed", 0)
        company = domain_context.get("company", "")
        scenario = domain_context.get("image_scenario", "")

        # Create deterministic variations
        import hashlib

        unique_string = f"{seed}_{company}_{scenario}_{professional_type}_{layout}"
        hash_value = int(hashlib.md5(unique_string.encode()).hexdigest()[:8], 16)

        # Use hash for deterministic but varied generation
        layout_variant = (hash_value // 5) % 4
        element_variant = (hash_value // 20) % 3

        if domain == "marketing_design":
            return self._create_marketing_visual(size, color_scheme, layout_variant, element_variant)
        elif domain == "scientific_publication":
            return self._create_scientific_diagram(size, color_scheme, layout_variant, element_variant)
        elif domain == "business_presentation":
            return self._create_business_presentation_image(size, color_scheme, layout_variant, element_variant)
        elif domain == "educational_content":
            return self._create_educational_material(size, color_scheme, layout_variant, element_variant)
        elif domain == "media_journalism":
            return self._create_media_graphic(size, color_scheme, layout_variant, element_variant)
        elif domain == "healthcare_communication":
            return self._create_healthcare_visual(size, color_scheme, layout_variant, element_variant)
        else:
            return self._create_business_presentation_image(size, color_scheme, layout_variant, element_variant)

    def _create_marketing_visual(self, size: Tuple[int, int], color_scheme: str, layout_variant: int, element_variant: int) -> Image.Image:
        """Create marketing-focused visual content."""
        img = Image.new("RGB", size, color="white")
        draw = ImageDraw.Draw(img)

        # Convert hex color to RGB
        primary_color = tuple(int(color_scheme[i : i + 2], 16) for i in (1, 3, 5))
        accent_color = tuple(min(255, c + 50) for c in primary_color)

        if layout_variant == 0:  # Hero banner style
            # Gradient background
            for y in range(size[1]):
                color_intensity = int(255 * (1 - y / size[1]))
                gradient_color = tuple(min(255, c + color_intensity // 3) for c in primary_color)
                draw.line([(0, y), (size[0], y)], fill=gradient_color)

            # Central focal element
            center_x, center_y = size[0] // 2, size[1] // 2
            draw.ellipse([center_x - 150, center_y - 100, center_x + 150, center_y + 100], fill=accent_color, outline="white", width=5)

            # Brand elements
            draw.rectangle([50, 50, 350, 150], fill=primary_color, outline="white", width=3)

        elif layout_variant == 1:  # Grid layout
            # Background
            draw.rectangle([0, 0, size[0], size[1]], fill=(248, 249, 250))

            # Grid elements
            grid_size = 4
            cell_w = size[0] // grid_size
            cell_h = size[1] // grid_size

            for i in range(grid_size):
                for j in range(grid_size):
                    if (i + j) % 2 == element_variant % 2:
                        x1, y1 = i * cell_w, j * cell_h
                        x2, y2 = x1 + cell_w - 10, y1 + cell_h - 10
                        draw.rectangle([x1 + 5, y1 + 5, x2, y2], fill=primary_color)

        elif layout_variant == 2:  # Diagonal composition
            # Dynamic diagonal background
            draw.polygon([(0, size[1]), (size[0] // 2, 0), (size[0], 0), (size[0], size[1])], fill=primary_color)
            draw.polygon([(0, 0), (0, size[1]), (size[0] // 2, 0)], fill=(240, 240, 240))

            # Accent elements
            for i in range(3):
                x = size[0] // 4 + i * size[0] // 6
                y = size[1] // 4 + i * size[1] // 8
                draw.ellipse([x - 40, y - 40, x + 40, y + 40], fill=accent_color)

        else:  # Modern minimal
            # Clean background with accent
            draw.rectangle([0, 0, size[0], size[1]], fill=(252, 252, 252))
            draw.rectangle([0, 0, size[0], size[1] // 6], fill=primary_color)

            # Geometric elements
            for i in range(3):
                x = size[0] // 4 + i * size[0] // 4
                y = size[1] // 2
                radius = 30 + i * 20
                draw.ellipse([x - radius, y - radius, x + radius, y + radius], fill=accent_color, outline=primary_color, width=3)

        return img

    def _create_scientific_diagram(self, size: Tuple[int, int], color_scheme: str, layout_variant: int, element_variant: int) -> Image.Image:
        """Create scientific/academic diagram."""
        img = Image.new("RGB", size, color="white")
        draw = ImageDraw.Draw(img)

        # Professional academic colors
        primary_color = (44, 62, 80)  # Dark blue-gray
        accent_color = (52, 152, 219)  # Professional blue

        if layout_variant == 0:  # Data visualization
            # Grid background
            grid_spacing = 50
            for x in range(0, size[0], grid_spacing):
                draw.line([(x, 0), (x, size[1])], fill=(230, 230, 230), width=1)
            for y in range(0, size[1], grid_spacing):
                draw.line([(0, y), (size[0], y)], fill=(230, 230, 230), width=1)

            # Data bars
            bar_width = 40
            data_points = [0.3, 0.7, 0.5, 0.9, 0.6, 0.8]
            for i, height_ratio in enumerate(data_points):
                x = 100 + i * 80
                bar_height = int(height_ratio * (size[1] - 200))
                y = size[1] - 100 - bar_height
                draw.rectangle([x, y, x + bar_width, size[1] - 100], fill=accent_color, outline=primary_color, width=2)

        elif layout_variant == 1:  # Process flow
            # Flow elements
            step_positions = [(150, 200), (400, 200), (650, 200), (400, 400)]
            for i, (x, y) in enumerate(step_positions):
                # Process boxes
                draw.rectangle([x - 60, y - 40, x + 60, y + 40], fill=accent_color, outline=primary_color, width=3)

                # Connecting arrows
                if i < len(step_positions) - 1:
                    next_x, next_y = step_positions[i + 1]
                    if i < 2:  # Horizontal arrows
                        draw.line([(x + 60, y), (next_x - 60, next_y)], fill=primary_color, width=4)
                        # Arrow head
                        draw.polygon([(next_x - 60, next_y), (next_x - 75, next_y - 10), (next_x - 75, next_y + 10)], fill=primary_color)
                    else:  # Vertical arrow
                        draw.line([(x, y + 40), (next_x, next_y - 40)], fill=primary_color, width=4)

        elif layout_variant == 2:  # Technical schematic
            # Border and title area
            draw.rectangle([20, 20, size[0] - 20, size[1] - 20], outline=primary_color, width=3)
            draw.rectangle([30, 30, size[0] - 30, 100], fill=(245, 245, 245))

            # Technical elements
            center_x, center_y = size[0] // 2, size[1] // 2

            # Central component
            draw.ellipse([center_x - 80, center_y - 80, center_x + 80, center_y + 80], fill=accent_color, outline=primary_color, width=4)

            # Surrounding components
            positions = [(center_x - 150, center_y), (center_x + 150, center_y), (center_x, center_y - 120), (center_x, center_y + 120)]
            for px, py in positions:
                draw.rectangle([px - 40, py - 25, px + 40, py + 25], fill=(220, 220, 220), outline=primary_color, width=2)
                # Connection lines
                draw.line([(center_x, center_y), (px, py)], fill=primary_color, width=2)

        else:  # Academic chart
            # Clean academic layout
            margin = 80
            chart_area = [margin, margin, size[0] - margin, size[1] - margin]
            draw.rectangle(chart_area, outline=primary_color, width=2)

            # Sample data visualization
            import math

            points = []
            for i in range(50):
                x = margin + (i / 49) * (size[0] - 2 * margin)
                y = margin + (size[1] - 2 * margin) * (0.5 + 0.3 * math.sin(i * 0.3))
                points.append((x, y))

            # Draw data line
            for i in range(len(points) - 1):
                draw.line([points[i], points[i + 1]], fill=accent_color, width=3)

        return img

    def _create_business_presentation_image(self, size: Tuple[int, int], color_scheme: str, layout_variant: int, element_variant: int) -> Image.Image:
        """Create business presentation visual."""
        img = Image.new("RGB", size, color="white")
        draw = ImageDraw.Draw(img)

        # Corporate colors
        primary_color = (31, 41, 55)  # Dark gray
        accent_color = (59, 130, 246)  # Professional blue

        if layout_variant == 0:  # Executive dashboard
            # Header
            draw.rectangle([0, 0, size[0], 80], fill=primary_color)

            # Metrics cards
            card_width = size[0] // 4 - 20
            for i in range(4):
                x = 10 + i * (card_width + 15)
                y = 100
                draw.rectangle([x, y, x + card_width, y + 120], fill=(248, 249, 250), outline=accent_color, width=2)

                # Metric visualization
                bar_height = 40 + (i * 15)
                draw.rectangle([x + 20, y + 80 - bar_height, x + card_width - 20, y + 80], fill=accent_color)

        elif layout_variant == 1:  # Strategic overview
            # Title section
            draw.rectangle([50, 50, size[0] - 50, 150], fill=(245, 245, 245), outline=primary_color, width=2)

            # Content sections
            sections = 3
            section_width = (size[0] - 120) // sections
            for i in range(sections):
                x = 60 + i * (section_width + 20)
                y = 200

                # Section box
                draw.rectangle([x, y, x + section_width, y + 200], fill=(252, 252, 252), outline=accent_color, width=2)

                # Content elements
                for j in range(3):
                    element_y = y + 30 + j * 50
                    draw.rectangle([x + 20, element_y, x + section_width - 20, element_y + 30], fill=accent_color)

        elif layout_variant == 2:  # Performance metrics
            # Chart background
            chart_x, chart_y = 100, 150
            chart_w, chart_h = size[0] - 200, size[1] - 300
            draw.rectangle([chart_x, chart_y, chart_x + chart_w, chart_y + chart_h], outline=primary_color, width=2)

            # Performance bars
            performance = [0.6, 0.8, 0.7, 0.9]
            bar_width = chart_w // 6

            for i, perf in enumerate(performance):
                x = chart_x + 50 + i * (bar_width + 30)
                bar_height = int(perf * (chart_h - 40))
                y = chart_y + chart_h - 20 - bar_height

                draw.rectangle([x, y, x + bar_width, chart_y + chart_h - 20], fill=accent_color, outline=primary_color, width=1)

        else:  # Corporate infographic
            # Background elements
            draw.rectangle([0, 0, size[0], size[1]], fill=(250, 250, 250))

            # Central focus area
            center_x, center_y = size[0] // 2, size[1] // 2
            draw.ellipse([center_x - 120, center_y - 80, center_x + 120, center_y + 80], fill=accent_color, outline=primary_color, width=4)

            # Surrounding information boxes
            box_positions = [(150, 150), (size[0] - 250, 150), (150, size[1] - 200), (size[0] - 250, size[1] - 200)]
            for bx, by in box_positions:
                draw.rectangle([bx, by, bx + 100, by + 60], fill=primary_color, outline=accent_color, width=2)

        return img

    def _create_educational_material(self, size: Tuple[int, int], color_scheme: str, layout_variant: int, element_variant: int) -> Image.Image:
        """Create educational/instructional visual."""
        img = Image.new("RGB", size, color="white")
        draw = ImageDraw.Draw(img)

        # Educational colors
        primary_color = (52, 152, 219)  # Friendly blue
        accent_color = (46, 204, 113)  # Success green

        if layout_variant == 0:  # Learning steps
            # Header
            draw.rectangle([0, 0, size[0], 80], fill=primary_color)

            # Step circles
            steps = 5
            step_spacing = (size[0] - 100) // (steps - 1)
            y_pos = 200

            for i in range(steps):
                x = 50 + i * step_spacing
                # Step circle
                draw.ellipse([x - 30, y_pos - 30, x + 30, y_pos + 30], fill=accent_color, outline=primary_color, width=3)

                # Connecting line
                if i < steps - 1:
                    next_x = 50 + (i + 1) * step_spacing
                    draw.line([(x + 30, y_pos), (next_x - 30, y_pos)], fill=primary_color, width=4)

        elif layout_variant == 1:  # Concept diagram
            # Central concept
            center_x, center_y = size[0] // 2, size[1] // 2
            draw.ellipse([center_x - 100, center_y - 60, center_x + 100, center_y + 60], fill=primary_color, outline=accent_color, width=4)

            # Related concepts
            concept_positions = [
                (center_x - 200, center_y - 150),
                (center_x + 200, center_y - 150),
                (center_x - 200, center_y + 150),
                (center_x + 200, center_y + 150),
            ]

            for cx, cy in concept_positions:
                draw.ellipse([cx - 60, cy - 40, cx + 60, cy + 40], fill=accent_color, outline=primary_color, width=2)
                # Connection lines
                draw.line([(center_x, center_y), (cx, cy)], fill=primary_color, width=2)

        elif layout_variant == 2:  # Instructional layout
            # Title area
            draw.rectangle([50, 50, size[0] - 50, 120], fill=(236, 240, 241), outline=primary_color, width=2)

            # Content blocks
            block_height = (size[1] - 200) // 3
            for i in range(3):
                y = 150 + i * (block_height + 20)

                # Number indicator
                draw.ellipse([70, y + 20, 110, y + 60], fill=accent_color)

                # Content area
                draw.rectangle([130, y, size[0] - 70, y + block_height], fill=(248, 249, 250), outline=primary_color, width=1)

        else:  # Knowledge visualization
            # Grid pattern for organized learning
            grid_cols, grid_rows = 4, 3
            cell_w = (size[0] - 100) // grid_cols
            cell_h = (size[1] - 150) // grid_rows

            for i in range(grid_rows):
                for j in range(grid_cols):
                    x = 50 + j * cell_w
                    y = 100 + i * cell_h

                    # Alternating pattern
                    if (i + j) % 2 == 0:
                        color = primary_color
                    else:
                        color = accent_color

                    draw.rectangle([x + 5, y + 5, x + cell_w - 5, y + cell_h - 5], fill=color, outline="white", width=2)

        return img

    def _create_media_graphic(self, size: Tuple[int, int], color_scheme: str, layout_variant: int, element_variant: int) -> Image.Image:
        """Create media/journalism visual."""
        img = Image.new("RGB", size, color="white")
        draw = ImageDraw.Draw(img)

        # Media colors
        primary_color = (231, 76, 60)  # News red
        accent_color = (52, 73, 94)  # Professional dark

        if layout_variant == 0:  # News layout
            # Header bar
            draw.rectangle([0, 0, size[0], 100], fill=primary_color)

            # News sections
            section_width = size[0] // 3
            for i in range(3):
                x = i * section_width
                y = 120

                # Section header
                draw.rectangle([x + 10, y, x + section_width - 10, y + 40], fill=accent_color)

                # Content blocks
                for j in range(3):
                    block_y = y + 60 + j * 80
                    draw.rectangle([x + 20, block_y, x + section_width - 20, block_y + 60], fill=(248, 249, 250), outline=accent_color, width=1)

        elif layout_variant == 1:  # Feature story
            # Large feature area
            draw.rectangle([50, 50, size[0] - 50, size[1] // 2], fill=(245, 245, 245), outline=primary_color, width=3)

            # Supporting elements
            support_y = size[1] // 2 + 30
            support_height = (size[1] - support_y - 50) // 2

            for i in range(2):
                y = support_y + i * (support_height + 20)
                draw.rectangle([70, y, size[0] - 70, y + support_height], fill=(252, 252, 252), outline=accent_color, width=2)

        elif layout_variant == 2:  # Editorial design
            # Diagonal layout
            points = [(0, size[1] // 3), (size[0] // 2, 0), (size[0], 0), (size[0], size[1]), (0, size[1])]
            draw.polygon(points, fill=primary_color)

            # Content overlay
            draw.rectangle([100, 150, size[0] - 100, size[1] - 150], fill=(255, 255, 255, 200), outline=accent_color, width=3)

        else:  # Information graphic
            # Timeline or process
            timeline_y = size[1] // 2
            points = 5
            point_spacing = (size[0] - 100) // (points - 1)

            # Timeline line
            draw.line([(50, timeline_y), (size[0] - 50, timeline_y)], fill=accent_color, width=6)

            for i in range(points):
                x = 50 + i * point_spacing
                # Timeline points
                draw.ellipse([x - 15, timeline_y - 15, x + 15, timeline_y + 15], fill=primary_color, outline=accent_color, width=3)

                # Information boxes
                box_y = timeline_y - 100 if i % 2 == 0 else timeline_y + 40
                draw.rectangle([x - 40, box_y, x + 40, box_y + 60], fill=(248, 249, 250), outline=primary_color, width=2)

        return img

    def _create_healthcare_visual(self, size: Tuple[int, int], color_scheme: str, layout_variant: int, element_variant: int) -> Image.Image:
        """Create healthcare/medical visual."""
        img = Image.new("RGB", size, color="white")
        draw = ImageDraw.Draw(img)

        # Healthcare colors
        primary_color = (39, 174, 96)  # Medical green
        accent_color = (41, 128, 185)  # Trust blue

        if layout_variant == 0:  # Medical diagram
            # Clean medical layout
            draw.rectangle([50, 50, size[0] - 50, size[1] - 50], outline=primary_color, width=3)

            # Central medical element
            center_x, center_y = size[0] // 2, size[1] // 2
            draw.ellipse([center_x - 80, center_y - 80, center_x + 80, center_y + 80], fill=primary_color, outline=accent_color, width=4)

            # Medical cross
            draw.rectangle([center_x - 10, center_y - 40, center_x + 10, center_y + 40], fill="white")
            draw.rectangle([center_x - 40, center_y - 10, center_x + 40, center_y + 10], fill="white")

        elif layout_variant == 1:  # Health information
            # Header with health theme
            draw.rectangle([0, 0, size[0], 80], fill=primary_color)

            # Information cards
            cards = 3
            card_width = (size[0] - 80) // cards
            for i in range(cards):
                x = 20 + i * (card_width + 20)
                y = 120

                draw.rectangle([x, y, x + card_width, y + 200], fill=(248, 249, 250), outline=accent_color, width=2)

                # Health icon area
                icon_center_x = x + card_width // 2
                icon_center_y = y + 60
                draw.ellipse([icon_center_x - 30, icon_center_y - 30, icon_center_x + 30, icon_center_y + 30], fill=primary_color)

        elif layout_variant == 2:  # Process flow
            # Health process steps
            steps = 4
            step_width = (size[0] - 100) // steps
            y_pos = size[1] // 2

            for i in range(steps):
                x = 50 + i * step_width + step_width // 2

                # Process step
                draw.rectangle([x - 40, y_pos - 50, x + 40, y_pos + 50], fill=accent_color, outline=primary_color, width=3)

                # Arrow to next step
                if i < steps - 1:
                    arrow_start_x = x + 40
                    arrow_end_x = 50 + (i + 1) * step_width + step_width // 2 - 40
                    draw.line([(arrow_start_x, y_pos), (arrow_end_x, y_pos)], fill=primary_color, width=4)
                    # Arrow head
                    draw.polygon([(arrow_end_x, y_pos), (arrow_end_x - 15, y_pos - 8), (arrow_end_x - 15, y_pos + 8)], fill=primary_color)

        else:  # Health education
            # Educational health layout
            margin = 60
            content_area = [margin, margin, size[0] - margin, size[1] - margin]
            draw.rectangle(content_area, outline=primary_color, width=2)

            # Health sections
            sections = 2
            section_height = (content_area[3] - content_area[1] - 40) // sections

            for i in range(sections):
                y = content_area[1] + 20 + i * (section_height + 20)

                # Section with health theme
                draw.rectangle([content_area[0] + 20, y, content_area[2] - 20, y + section_height], fill=(245, 255, 245), outline=accent_color, width=2)

                # Health indicator
                indicator_x = content_area[0] + 40
                indicator_y = y + section_height // 2
                draw.ellipse([indicator_x - 15, indicator_y - 15, indicator_x + 15, indicator_y + 15], fill=primary_color)

        return img

    def _create_enhanced_image_by_type(self, image_type: str, size: Tuple[int, int]) -> Image.Image:
        """Create enhanced image by type for backward compatibility."""
        if image_type == "marketing_visual":
            return self._create_marketing_visual(size, "#3498db", 0, 0)
        elif image_type == "scientific_diagram":
            return self._create_scientific_diagram(size, "#2c3e50", 0, 0)
        elif image_type == "educational_material":
            return self._create_educational_material(size, "#9b59b6", 0, 0)
        else:  # corporate_presentation or default
            return self._create_business_presentation_image(size, "#34495e", 0, 0)

    def _create_geometric_image(self, size):
        """Create geometric pattern image."""
        img = Image.new("RGB", size, color="white")
        draw = ImageDraw.Draw(img)

        # Draw colorful geometric shapes
        colors = [(255, 0, 0), (0, 255, 0), (0, 0, 255), (255, 255, 0), (255, 0, 255)]

        for i in range(3):
            color = colors[i % len(colors)]
            x1, y1 = i * 150, i * 100
            x2, y2 = x1 + 200, y1 + 150
            draw.rectangle([x1, y1, x2, y2], fill=color, outline="black", width=3)

        # Add some circles
        for i in range(2):
            color = colors[(i + 3) % len(colors)]
            x, y = 100 + i * 200, 50 + i * 100
            draw.ellipse([x, y, x + 100, y + 100], fill=color, outline="black", width=2)

        return img

    def _create_colorful_image(self, size):
        """Create colorful image for filter effects."""
        img = Image.new("RGB", size, color="white")
        draw = ImageDraw.Draw(img)

        # Create colorful background gradient
        for x in range(size[0]):
            for y in range(size[1]):
                r = int(255 * x / size[0])
                g = int(255 * y / size[1])
                b = int(255 * (x + y) / (size[0] + size[1]))
                img.putpixel((x, y), (r % 256, g % 256, b % 256))

        # Add some shapes on top
        colors = [(255, 255, 255), (0, 0, 0), (255, 0, 0), (0, 255, 255)]
        for i in range(4):
            color = colors[i]
            x, y = 50 + i * 100, 50 + i * 50
            draw.ellipse([x, y, x + 80, y + 80], fill=color, outline="black")

        return img

    def _create_text_image(self, size):
        """Create text-based image."""
        img = Image.new("RGB", size, color="lightblue")
        draw = ImageDraw.Draw(img)

        # Try to use a font, fallback to default if not available
        try:
            font = ImageFont.truetype("arial.ttf", 48)
        except Exception:
            font = ImageFont.load_default()

        # Add text content
        texts = ["SAMPLE", "IMAGE", "FOR", "EDITING"]
        for i, text in enumerate(texts):
            x = size[0] // 4
            y = 50 + i * 80
            draw.text((x, y), text, fill="black", font=font)

        # Add decorative border
        draw.rectangle([10, 10, size[0] - 10, size[1] - 10], outline="red", width=5)

        return img

    def _create_gradient_image(self, size):
        """Create gradient image."""
        img = Image.new("RGB", size, color="white")

        # Create horizontal gradient
        for x in range(size[0]):
            color_val = int(255 * x / size[0])
            for y in range(size[1]):
                img.putpixel((x, y), (color_val, 100, 255 - color_val))

        return img

    def _generate_document_template(self, task_data: Dict[str, Any], temp_dir: str) -> str:
        """Generate professional LibreOffice Writer document template."""
        template_type = task_data.get("document_template_type", "empty")

        # Get professional context
        professional_scenario = task_data.get("professional_scenario", {})
        domain_context = task_data.get("domain_context", {})

        doc = Document()

        if template_type == "professional_empty" or template_type == "empty":
            # Professional empty document with header
            if professional_scenario:
                company = domain_context.get("company", "Professional Company")
                project_type = domain_context.get("project_type", "Image Processing Project")

                # Professional header
                doc.add_heading(f"{company}", 0)
                doc.add_heading(f"{project_type.replace('_', ' ').title()}", level=1)
                doc.add_paragraph("")  # Space for content
            else:
                # Fallback to basic empty document
                pass

        elif template_type == "level2_with_content":
            # Level 2: Brief content for resize + insertion tasks
            brief_content = self._generate_brief_content(professional_scenario, domain_context, "level2")
            doc.add_paragraph(brief_content)
            doc.add_paragraph("")  # Space for image insertion

        elif template_type == "level3_with_content":
            # Level 3: Brief content for modify + caption tasks
            brief_content = self._generate_brief_content(professional_scenario, domain_context, "level3")
            doc.add_paragraph(brief_content)
            doc.add_paragraph("")  # Space for image and caption

        elif template_type == "professional_layout" or template_type == "with_placeholder":
            # Professional layout template
            if professional_scenario:
                company = domain_context.get("company", "Professional Company")
                project_type = domain_context.get("project_type", "Image Processing Project")
                context = domain_context.get("context", "professional documentation")

                doc.add_heading(f"{company}", 0)
                doc.add_heading(f"{project_type.replace('_', ' ').title()} Documentation", level=1)
                doc.add_paragraph("")

                doc.add_heading("Project Overview", level=2)
                doc.add_paragraph(f"This document contains visual assets and documentation for {context}.")
                doc.add_paragraph("")

                doc.add_heading("Image Processing", level=2)
                doc.add_paragraph("Please insert the processed image below:")
                doc.add_paragraph("")  # Space for image
                doc.add_paragraph("Image specifications and quality have been optimized for professional use.")
            else:
                # Fallback to basic template
                doc.add_paragraph("Image Analysis Report")
                doc.add_paragraph("")
                doc.add_paragraph("Insert the resized image below:")
                doc.add_paragraph("")
                doc.add_paragraph("The image has been processed according to the specified dimensions.")

        elif template_type == "professional_caption_layout" or template_type == "with_caption_area":
            # Professional caption template
            if professional_scenario:
                company = domain_context.get("company", "Professional Company")
                project_type = domain_context.get("project_type", "Image Processing Project")
                context = domain_context.get("context", "professional documentation")
                domain = domain_context.get("domain", "business")

                doc.add_heading(f"{company}", 0)
                doc.add_heading(f"{project_type.replace('_', ' ').title()} - Enhanced Content", level=1)
                doc.add_paragraph("")

                doc.add_heading("Project Summary", level=2)
                doc.add_paragraph(f"Visual content processing for {context} in the {domain.replace('_', ' ')} domain.")
                doc.add_paragraph("")

                doc.add_heading("Enhanced Image Processing", level=2)
                doc.add_paragraph("The following image has been processed with professional filters and enhancements:")
                doc.add_paragraph("")  # Space for image and caption
                doc.add_paragraph("")  # Additional space for caption

                doc.add_heading("Technical Notes", level=2)
                doc.add_paragraph("All visual processing has been completed according to professional standards.")
            else:
                # Fallback to basic template
                doc.add_paragraph("Image Processing Analysis")
                doc.add_paragraph("")
                doc.add_paragraph("Filtered Image Results:")
                doc.add_paragraph("Insert the filtered image below with appropriate caption:")
                doc.add_paragraph("")

        # Save document
        document_filename = task_data.get("document_file", "document.docx")
        document_path = os.path.join(temp_dir, document_filename)
        doc.save(document_path)

        return document_path

    def _generate_brief_content(self, professional_scenario: Dict[str, Any], domain_context: Dict[str, Any], level: str) -> str:
        """Generate brief content (20-30 words) using combinatorial pool."""
        domain = domain_context.get("domain", "business_presentation")
        company = domain_context.get("company", "Professional Company")
        project_type = domain_context.get("project_type", "project")
        context = domain_context.get("context", "documentation")

        # Get company abbreviation for brevity
        company_short = company.split()[0]

        # Domain-specific brief content templates (20-30 words each)
        content_templates = {
            "marketing_design": [
                f"{company_short} marketing campaign overview. Target audience engagement through strategic visual content.",
                f"Brand identity guidelines for {company_short}. Visual consistency across all marketing materials required.",
                f"Creative brief: {project_type}. Brand messaging and visual elements for promotional activities.",
                f"Marketing strategy document. {company_short} brand positioning and visual communication framework.",
                f"Campaign planning document. Visual content requirements for {context} implementation.",
            ],
            "scientific_publication": [
                f"{company_short} research findings summary. Data visualization and scientific documentation required.",
                f"Laboratory report: {project_type}. Experimental results and visual evidence documentation.",
                f"Research methodology overview. {company_short} study protocols and analysis procedures.",
                "Scientific publication draft. Research data presentation and academic documentation standards.",
                f"Study documentation: {context}. Data collection and visual analysis requirements.",
            ],
            "business_presentation": [
                f"{company_short} quarterly report. Business metrics and performance analysis documentation.",
                f"Executive summary: {project_type}. Strategic planning and operational overview presentation.",
                f"Business proposal document. {company_short} strategic initiatives and implementation plan.",
                "Corporate presentation materials. Financial data and business intelligence visualization.",
                f"Strategic planning document: {context}. Business objectives and performance metrics.",
            ],
            "educational_content": [
                f"{company_short} curriculum development. Educational materials and instructional design documentation.",
                f"Learning module: {project_type}. Educational content and visual learning aids.",
                f"Course documentation. {company_short} instructional methodology and learning outcomes.",
                "Educational resource guide. Learning materials and pedagogical content organization.",
                f"Academic documentation: {context}. Educational standards and instructional content delivery.",
            ],
            "media_journalism": [
                f"{company_short} editorial content. News reporting and journalistic documentation standards.",
                f"Media coverage: {project_type}. Editorial guidelines and content publication requirements.",
                f"Press documentation. {company_short} journalism standards and editorial processes.",
                "News report outline. Editorial content and media publication framework.",
                f"Journalistic documentation: {context}. Editorial standards and media content guidelines.",
            ],
            "healthcare_communication": [
                f"{company_short} health information guide. Medical documentation and patient communication standards.",
                f"Healthcare documentation: {project_type}. Medical information and patient education materials.",
                f"Health communication guidelines. {company_short} medical content and patient information.",
                "Medical documentation standards. Healthcare information and clinical communication protocols.",
                f"Patient information: {context}. Health education and medical documentation requirements.",
            ],
        }

        # Select appropriate template based on domain
        templates = content_templates.get(domain, content_templates["business_presentation"])
        return templates[hash(str(professional_scenario)) % len(templates)]

    def _create_expected_files(self, task_data: Dict[str, Any], task_id: str, temp_dir: str, source_image_path: str, document_template_path: str) -> Dict[str, str]:
        """Create expected/ground truth files for evaluation using the ACTUAL source files."""
        expected_files = {}
        task_type = task_data.get("task_type")

        if task_type == "image_format_conversion":
            # Create expected converted image using the ACTUAL source image
            expected_image = self._create_expected_converted_image(task_data, temp_dir, source_image_path)

            # Create expected document with the converted image
            expected_doc = self._create_expected_document_with_image(task_data, temp_dir, expected_image, document_template_path)
            # Provide both keys:
            # - expected_document for S3 upload and evaluator download
            # - gold_standard_file for local ground_truth packaging
            expected_files["expected_document"] = expected_doc
            expected_files["gold_standard_file"] = expected_doc

        elif task_type == "image_resizing_placement":
            # Create expected resized image using the ACTUAL source image
            expected_image = self._create_expected_resized_image(task_data, temp_dir, source_image_path)

            # Create expected document with positioned image
            expected_doc = self._create_expected_document_with_positioned_image(task_data, temp_dir, expected_image, document_template_path)
            expected_files["expected_document"] = expected_doc
            expected_files["gold_standard_file"] = expected_doc

        elif task_type == "image_filter_caption":
            # Create expected filtered image using the ACTUAL source image
            expected_image = self._create_expected_filtered_image(task_data, temp_dir, source_image_path)

            # Create expected document with image and caption
            expected_doc = self._create_expected_document_with_caption(task_data, temp_dir, expected_image, document_template_path)
            expected_files["expected_document"] = expected_doc
            expected_files["gold_standard_file"] = expected_doc

        elif task_type == "basic_image_insertion":
            # Create expected document with inserted image (Level 1)
            expected_doc = self._create_expected_document_with_image(task_data, temp_dir, source_image_path, document_template_path)
            expected_files["expected_document"] = expected_doc
            expected_files["gold_standard_file"] = expected_doc

        elif task_type == "image_resize_insertion":
            # Create expected resized image using the ACTUAL source image (Level 2)
            expected_image = self._create_expected_resized_image(task_data, temp_dir, source_image_path)

            # Create expected document with resized image
            expected_doc = self._create_expected_document_with_image(task_data, temp_dir, expected_image, document_template_path)
            expected_files["expected_document"] = expected_doc
            expected_files["gold_standard_file"] = expected_doc

        elif task_type == "image_modify_caption":
            # Create expected modified image (resize OR grayscale) (Level 3)
            modification_type = task_data.get("modification_type", "resize")
            if modification_type == "resize":
                expected_image = self._create_expected_resized_image(task_data, temp_dir, source_image_path)
            else:  # grayscale
                expected_image = self._create_expected_grayscale_image(task_data, temp_dir, source_image_path)

            # Create expected document with image and caption
            expected_doc = self._create_expected_document_with_caption(task_data, temp_dir, expected_image, document_template_path)
            expected_files["expected_document"] = expected_doc
            expected_files["gold_standard_file"] = expected_doc

            # Create JSON specification for document evaluation
            json_spec = self._create_json_specification(task_data, temp_dir)
            expected_files["json_specification"] = json_spec

        return expected_files

    def _create_expected_converted_image(self, task_data: Dict[str, Any], temp_dir: str, source_image_path: str) -> str:
        """Create expected converted image for Level 1 using the ACTUAL source image."""
        # Load the actual source image that was generated for this task
        try:
            img = Image.open(source_image_path)
        except Exception as e:
            print(f"Warning: Could not load source image {source_image_path}: {e}")
            # Fallback to a basic image if source can't be loaded
            img = Image.new("RGB", (800, 600), color="white")

        # Save in target format with expected filename
        target_filename = task_data.get("target_image", "converted.jpg")
        expected_path = os.path.join(temp_dir, f"expected_{target_filename}")

        if target_filename.lower().endswith(".jpg") or target_filename.lower().endswith(".jpeg"):
            img.save(expected_path, "JPEG", quality=95)
        elif target_filename.lower().endswith(".png"):
            img.save(expected_path, "PNG")
        elif target_filename.lower().endswith(".gif"):
            img.save(expected_path, "GIF")
        else:
            img.save(expected_path, "PNG")

        return expected_path

    def _create_expected_resized_image(self, task_data: Dict[str, Any], temp_dir: str, source_image_path: str) -> str:
        """Create expected resized image for Level 2 using the ACTUAL source image."""
        # Load the actual source image that was generated for this task
        try:
            img = Image.open(source_image_path)
        except Exception as e:
            print(f"Warning: Could not load source image {source_image_path}: {e}")
            # Fallback to a basic image if source can't be loaded
            img = Image.new("RGB", (1600, 1200), color="white")

        # Resize to target dimensions from task data
        target_width = task_data.get("target_width", 200)
        target_height = task_data.get("target_height", 150)
        resized_img = img.resize((target_width, target_height), Image.Resampling.LANCZOS)

        # Save resized image with expected filename
        resized_filename = task_data.get("resized_image", "resized.png")
        expected_path = os.path.join(temp_dir, f"expected_{resized_filename}")
        resized_img.save(expected_path, "PNG")

        return expected_path

    def _create_expected_filtered_image(self, task_data: Dict[str, Any], temp_dir: str, source_image_path: str) -> str:
        """Create expected filtered image for Level 3 using the ACTUAL source image."""
        # Load the actual source image that was generated for this task
        try:
            img = Image.open(source_image_path)
        except Exception as e:
            print(f"Warning: Could not load source image {source_image_path}: {e}")
            # Fallback to a basic image if source can't be loaded
            img = Image.new("RGB", (1200, 800), color="white")

        # Apply filter based on filter type from task data
        filter_type = task_data.get("filter_type", "Grayscale")

        if filter_type == "Grayscale":
            filtered_img = img.convert("L").convert("RGB")
        elif filter_type == "Sepia":
            filtered_img = self._apply_sepia_filter(img.copy())  # Make a copy to avoid modifying original
        elif filter_type == "Invert Colors":
            filtered_img = self._apply_invert_filter(img.copy())  # Make a copy
        elif filter_type == "Blur":
            from PIL import ImageFilter

            filtered_img = img.filter(ImageFilter.BLUR)
        else:
            filtered_img = img.convert("L").convert("RGB")  # Default to grayscale

        # Save filtered image with expected filename
        filtered_filename = task_data.get("filtered_image", "filtered.png")
        expected_path = os.path.join(temp_dir, f"expected_{filtered_filename}")
        filtered_img.save(expected_path, "PNG")

        return expected_path

    def _apply_sepia_filter(self, img):
        """Apply sepia filter to image."""
        pixels = img.load()
        width, height = img.size

        for x in range(width):
            for y in range(height):
                r, g, b = pixels[x, y]

                # Sepia formula
                tr = int(0.393 * r + 0.769 * g + 0.189 * b)
                tg = int(0.349 * r + 0.686 * g + 0.168 * b)
                tb = int(0.272 * r + 0.534 * g + 0.131 * b)

                # Clamp values
                tr = min(255, tr)
                tg = min(255, tg)
                tb = min(255, tb)

                pixels[x, y] = (tr, tg, tb)

        return img

    def _apply_invert_filter(self, img):
        """Apply color inversion filter."""
        pixels = img.load()
        width, height = img.size

        for x in range(width):
            for y in range(height):
                r, g, b = pixels[x, y]
                pixels[x, y] = (255 - r, 255 - g, 255 - b)

        return img

    def _create_expected_grayscale_image(self, task_data: Dict[str, Any], temp_dir: str, source_image_path: str) -> str:
        """Create expected grayscale image for Level 3 using the ACTUAL source image."""
        # Load the actual source image that was generated for this task
        try:
            img = Image.open(source_image_path)
        except Exception as e:
            print(f"Warning: Could not load source image {source_image_path}: {e}")
            # Fallback to a basic image if source can't be loaded
            img = Image.new("RGB", (1200, 900), color="white")

        # Convert to grayscale
        grayscale_img = img.convert("L").convert("RGB")

        # Save grayscale image with expected filename
        modified_filename = task_data.get("modified_image", "grayscale.png")
        expected_path = os.path.join(temp_dir, f"expected_{modified_filename}")
        grayscale_img.save(expected_path, "PNG")

        return expected_path

    def _create_expected_document_with_image(self, task_data: Dict[str, Any], temp_dir: str, image_path: str, document_template_path: str) -> str:
        """Create expected document with image for Level 1 using the ACTUAL document template."""
        # Load the actual document template that was generated for this task
        try:
            doc = Document(document_template_path)
        except Exception as e:
            print(f"Warning: Could not load document template {document_template_path}: {e}")
            # Fallback to a new document
            doc = Document()
            doc.add_paragraph("Document with converted image:")

        # Add the converted image to the document
        try:
            doc.add_picture(image_path, width=Inches(4), height=Inches(3))
        except Exception as e:
            print(f"Warning: Could not add image to document: {e}")
            doc.add_paragraph("[Image could not be loaded]")

        # Save expected document with proper filename
        document_filename = task_data.get("document_file", "document.docx")
        expected_path = os.path.join(temp_dir, f"expected_{document_filename}")
        doc.save(expected_path)

        return expected_path

    def _create_expected_document_with_positioned_image(self, task_data: Dict[str, Any], temp_dir: str, image_path: str, document_template_path: str) -> str:
        """Create expected document with positioned image for Level 2 using the ACTUAL document template."""
        # Load the actual document template that was generated for this task
        try:
            doc = Document(document_template_path)
        except Exception as e:
            print(f"Warning: Could not load document template {document_template_path}: {e}")
            # Fallback to a new document with basic content
            doc = Document()
            doc.add_paragraph("Image Analysis Report")
            doc.add_paragraph("")
            doc.add_paragraph("Insert the resized image below:")

        # Add the resized image with professional sizing
        try:
            doc.add_picture(image_path, width=Inches(4), height=Inches(3))
        except Exception as e:
            print(f"Warning: Could not add image to document: {e}")
            doc.add_paragraph("[Resized image could not be loaded]")

        # Add professional conclusion
        doc.add_paragraph("The image has been processed according to the specified dimensions.")

        # Save expected document with proper filename
        document_filename = task_data.get("document_file", "document_template.docx")
        expected_path = os.path.join(temp_dir, f"expected_{document_filename}")
        doc.save(expected_path)

        return expected_path

    def _create_expected_document_with_caption(self, task_data: Dict[str, Any], temp_dir: str, image_path: str, document_template_path: str) -> str:
        """Create expected document with image and caption for Level 3 using the ACTUAL document template."""
        # Load the actual document template that was generated for this task
        try:
            doc = Document(document_template_path)
        except Exception as e:
            print(f"Warning: Could not load document template {document_template_path}: {e}")
            # Fallback to a new document with basic content
            doc = Document()
            doc.add_paragraph("Image Processing Analysis")
            doc.add_paragraph("")
            doc.add_paragraph("Filtered Image Results:")
            doc.add_paragraph("Insert the filtered image below with appropriate caption:")

        # Add the filtered image with professional sizing
        try:
            doc.add_picture(image_path, width=Inches(4), height=Inches(3))
        except Exception as e:
            print(f"Warning: Could not add image to document: {e}")
            doc.add_paragraph("[Filtered image could not be loaded]")

        # Add the professional caption from task data
        caption_text = task_data.get("caption_text", "Filtered image")
        doc.add_paragraph(caption_text)

        # Save expected document with proper filename
        document_filename = task_data.get("document_file", "document_template.docx")
        expected_path = os.path.join(temp_dir, f"expected_{document_filename}")
        doc.save(expected_path)

        return expected_path

    def _create_json_specification(self, task_data: Dict[str, Any], temp_dir: str) -> str:
        """Create JSON specification file for level 3 document evaluation."""
        # Extract task data for specification generation
        caption_text = task_data.get("caption_text", "")
        domain_context = task_data.get("domain_context", {})

        # Generate brief content that would be in the document
        professional_scenario = task_data.get("professional_scenario", {})
        brief_content = self._generate_brief_content(professional_scenario, domain_context, "level3")

        # Create JSON specification containing only global requirements
        specification = {
            "global": {
                "required_texts": [
                    brief_content.strip(),  # Original document content
                    caption_text.strip(),  # Caption text to be added
                ],
                "required_headings": [],
            }
        }

        # Save JSON specification file
        document_filename = task_data.get("document_file", "document.docx")
        base_name = document_filename.replace(".docx", "")
        spec_filename = f"{base_name}_spec.json"
        spec_path = os.path.join(temp_dir, spec_filename)

        with open(spec_path, "w", encoding="utf-8") as f:
            json.dump(specification, f, indent=2, ensure_ascii=False)

        return spec_path


class ImageProcessingConfigProvider(ConfigProviderInterface):
    """Config provider implementation for image processing tasks."""

    def __init__(self):
        """Initialize image processing config provider."""
        # Image processing specific task types
        self.supported_tasks = {
            "basic_image_insertion",
            "image_resize_insertion",
            "image_modify_caption",
        }

        # Evaluation mode mapping
        self.evaluation_mode_mapping = {
            "basic_image_insertion": "multi_evaluator",
            "image_resize_insertion": "multi_evaluator",
            "image_modify_caption": "multi_evaluator",
        }

    def build_setup_steps(
        self,
        task_data: Dict[str, Any],
        s3_urls: Dict[str, str],
        files_created: Dict[str, str],
    ) -> List[Dict[str, Any]]:
        """Build setup configuration steps for image processing tasks."""
        steps = []
        task_type = task_data.get("task_type")

        if not self.supports_task_type(task_type):
            return steps

        # Step 1: Download files (no need to create Desktop directory - exists in Ubuntu)
        download_files = []

        # Download source image
        if "source_image" in s3_urls:
            source_image_path = f"/home/user/Desktop/{task_data.get('source_image', 'source_image.png')}"
            download_files.append({"url": s3_urls["source_image"], "path": source_image_path})

        # Download document template
        if "document_template" in s3_urls:
            document_path = f"/home/user/Desktop/{task_data.get('document_file', 'document.docx')}"
            download_files.append({"url": s3_urls["document_template"], "path": document_path})

        if download_files:
            steps.append({"type": "download", "parameters": {"files": download_files}})

        # Step 2: Launch applications based on task type
        evaluation_data = task_data.get("evaluation_data", {})
        requires_gimp = evaluation_data.get("requires_gimp", True)

        # Launch GIMP only if required (not for Level 1)
        if requires_gimp:
            source_image_path = f"/home/user/Desktop/{task_data.get('source_image', 'source_image.png')}"
            steps.append({"type": "launch", "parameters": {"command": ["gimp", source_image_path]}})

        # Step 3: Launch LibreOffice Writer with document (always required)
        document_path = f"/home/user/Desktop/{task_data.get('document_file', 'document.docx')}"
        steps.append(
            {
                "type": "launch",
                "parameters": {"command": ["libreoffice", "--writer", document_path]},
            }
        )

        # Add sleep period at the end of setup
        steps.append({"type": "sleep", "parameters": {"seconds": 10.0}})

        return steps

    def get_evaluation_mode(self, task_type: str, level: int) -> str:
        """Get evaluation mode for task type and level."""
        return self.evaluation_mode_mapping.get(task_type, "multi_evaluator")

    def supports_task_type(self, task_type: str) -> bool:
        """Check if this provider supports the given task type."""
        return task_type in self.supported_tasks


class ImageProcessingEvaluationProvider(EvaluationProviderInterface):
    """Evaluation provider implementation for image processing tasks."""

    def __init__(self):
        """Initialize image processing evaluation provider."""
        # Image processing specific task types
        self.supported_tasks = {
            "basic_image_insertion",
            "image_resize_insertion",
            "image_modify_caption",
        }

        # Initialize evaluators
        self.evaluators = ImageProcessingEvaluators()

    def build_evaluator_config(
        self,
        task_data: Dict[str, Any],
        files_created: Dict[str, str],
        evaluation_mode: str,
        s3_urls: Dict[str, str],
    ) -> Dict[str, Any]:
        """Build evaluator configuration for the task."""
        task_type = task_data.get("task_type")

        if not self.supports_task_type(task_type):
            return {}

        # Delegate to the original evaluators which have the proper structure
        # Add evaluation_mode to task_data for compatibility
        task_data_with_mode = {**task_data, "evaluation_mode": evaluation_mode}

        # All image processing tasks use multi-evaluator
        return self.evaluators.build_multi_evaluator_config(task_type, task_data_with_mode, files_created, s3_urls)

    def supports_task_type(self, task_type: str) -> bool:
        """Check if this provider supports the given task type."""
        return task_type in self.supported_tasks

    def get_evaluator_instance(self):
        """Get the evaluator instance for this category."""
        return self.evaluators

    def get_supported_task_types(self) -> set:
        """Get set of supported task types."""
        return self.supported_tasks.copy()


__all__ = [
    "ImageProcessingFileProvider",
    "ImageProcessingConfigProvider",
    "ImageProcessingEvaluationProvider",
]
